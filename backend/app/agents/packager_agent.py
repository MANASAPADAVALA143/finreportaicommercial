"""PACKAGER — assembles XLSX + DOCX + PDF from DB statements, notes, and commentary (pure Python)."""
from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace
from typing import Any

from fpdf import FPDF
from openpyxl import Workbook
from sqlalchemy.orm import Session

from app.agents.base_agent import AgentContext, BaseAgent
from app.models.ifrs_statement import DisclosureNote, GeneratedStatement, IFRSStatementKind, StatementLineItem, TrialBalance


EXPORT_ROOT = Path(__file__).resolve().parents[2] / "agentic_exports"


def _lines_for_statement(db: Session, trial_balance_id: int, tenant_id: str, kind: IFRSStatementKind) -> list[StatementLineItem]:
    stmt = (
        db.query(GeneratedStatement)
        .filter(
            GeneratedStatement.trial_balance_id == trial_balance_id,
            GeneratedStatement.tenant_id == tenant_id,
            GeneratedStatement.statement_type == kind,
        )
        .first()
    )
    if not stmt:
        return []
    return (
        db.query(StatementLineItem)
        .filter(StatementLineItem.statement_id == stmt.id)
        .order_by(StatementLineItem.display_order)
        .all()
    )


def _sheet_title(kind: IFRSStatementKind) -> str:
    return {
        IFRSStatementKind.financial_position: "BS",
        IFRSStatementKind.profit_loss: "PL",
        IFRSStatementKind.cash_flows: "CF",
        IFRSStatementKind.equity: "SOCE",
    }.get(kind, kind.value[:8])


def _lines_from_vault_payload(st: dict[str, Any] | None, kind: IFRSStatementKind) -> list[Any]:
    if not st:
        return []
    lines = st.get(kind.value) or []
    out: list[Any] = []
    for item in lines:
        if not isinstance(item, dict):
            continue
        out.append(
            SimpleNamespace(
                ifrs_section="",
                ifrs_line_item=str(item.get("line") or ""),
                amount=item.get("amount") or 0,
            )
        )
    return out


def _comparative_sheet_rows(cur_lines: list, prior_lines: list | None) -> list[list[Any]]:
    prior_map: dict[tuple[str, str], float] = {}
    for li in prior_lines or []:
        sec = getattr(li, "ifrs_section", "") or ""
        name = (getattr(li, "ifrs_line_item", None) or "").strip().casefold()
        prior_map[(sec, name)] = float(getattr(li, "amount", 0) or 0)
    rows: list[list[Any]] = []
    for li in cur_lines:
        sec = getattr(li, "ifrs_section", "") or ""
        name = (getattr(li, "ifrs_line_item", None) or "").strip().casefold()
        cur_amt = float(getattr(li, "amount", 0) or 0)
        rows.append([sec, getattr(li, "ifrs_line_item", ""), cur_amt, prior_map.get((sec, name), "")])
    return rows


def _year_labels(ctx: AgentContext, tb) -> tuple[str, str]:
    cy = str(tb.period_end.year) if tb.period_end else "Current"
    py = str(tb.period_end.year - 1) if tb.period_end else "Prior"
    pid = ctx.shared.get("prior_trial_balance_id")
    if pid:
        ptb = ctx.db.query(TrialBalance).filter(TrialBalance.id == int(pid)).first()
        if ptb and ptb.period_end:
            py = str(ptb.period_end.year)
    return cy, py


class PackagerAgent(BaseAgent):
    agent_id = "PACKAGER"

    def run(self) -> dict[str, str]:
        ctx = self.ctx
        tb = ctx.db.query(TrialBalance).filter(TrialBalance.id == ctx.trial_balance_id).first()
        if not tb:
            raise RuntimeError("Trial balance not found")
        company = (tb.company_name or "Company").replace("/", "-")[:80]
        period = (tb.period_end.isoformat() if tb.period_end else "Period")[:32]

        out_dir = EXPORT_ROOT / ctx.public_run_id
        out_dir.mkdir(parents=True, exist_ok=True)
        base = f"{company}_IFRS_{period}"
        xlsx_path = out_dir / f"{base}_Statements.xlsx"
        docx_path = out_dir / f"{base}_Notes.docx"
        pdf_path = out_dir / f"{base}_Pack.pdf"

        # --- Excel (IAS 1 comparative columns when prior TB or vault snapshot exists) ---
        wb = Workbook()
        wb.remove(wb.active)
        cy_l, py_l = _year_labels(ctx, tb)
        pid = ctx.shared.get("prior_trial_balance_id")
        vault_st = (ctx.shared.get("prior_vault_snapshot") or {}).get("statements")
        comparative = bool(pid) or bool(vault_st)
        for kind in (
            IFRSStatementKind.financial_position,
            IFRSStatementKind.profit_loss,
            IFRSStatementKind.cash_flows,
            IFRSStatementKind.equity,
        ):
            ws = wb.create_sheet(title=_sheet_title(kind)[:31])
            lines = _lines_for_statement(ctx.db, ctx.trial_balance_id, ctx.tenant_id, kind)
            if comparative:
                prior_lines: list | None = None
                if pid:
                    prior_lines = _lines_for_statement(ctx.db, int(pid), ctx.tenant_id, kind)
                elif vault_st:
                    prior_lines = _lines_from_vault_payload(vault_st, kind)
                ws.append(["Section", "Line item", f"FY {cy_l}", f"FY {py_l} (prior)"])
                for row in _comparative_sheet_rows(lines, prior_lines):
                    ws.append(row)
            else:
                ws.append(["Section", "Line item", "Amount"])
                for li in lines:
                    ws.append([li.ifrs_section, li.ifrs_line_item, float(li.amount or 0)])
        wb.save(str(xlsx_path))

        # --- Notes DOCX (optional dependency) ---
        notes = (
            ctx.db.query(DisclosureNote)
            .filter(
                DisclosureNote.trial_balance_id == ctx.trial_balance_id,
                DisclosureNote.tenant_id == ctx.tenant_id,
            )
            .order_by(DisclosureNote.note_number)
            .all()
        )
        try:
            from docx import Document

            doc = Document()
            doc.add_heading(f"{company} — IFRS Notes", level=0)
            for n in notes:
                doc.add_heading(f"{n.note_code} {n.note_title}", level=1)
                doc.add_paragraph((n.user_edited_content or n.ai_generated_content or "")[:8000])
            doc.save(str(docx_path))
        except Exception:
            # Fallback: minimal RTF-like text file with .docx extension avoided — write plain text
            txt_fallback = out_dir / f"{base}_Notes.txt"
            parts = [f"{n.note_code} {n.note_title}\n{(n.user_edited_content or n.ai_generated_content or '')}\n" for n in notes]
            txt_fallback.write_text("\n".join(parts), encoding="utf-8")
            docx_path = txt_fallback

        # --- PDF pack ---
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 10, f"{company} — IFRS Pack", ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 8, f"Period: {period}", ln=True)
        comm = ctx.shared.get("commentary") or {}
        if isinstance(comm, dict):
            pdf.ln(4)
            for k, v in comm.items():
                pdf.set_font("Helvetica", "B", 11)
                pdf.multi_cell(0, 7, str(k).replace("_", " ").title())
                pdf.set_font("Helvetica", "", 10)
                pdf.multi_cell(0, 6, str(v)[:3500])
                pdf.ln(2)
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Statement highlights (P&L)", ln=True)
        pdf.set_font("Helvetica", "", 8)
        pl_lines = _lines_for_statement(ctx.db, ctx.trial_balance_id, ctx.tenant_id, IFRSStatementKind.profit_loss)
        cy_l, py_l = _year_labels(ctx, tb)
        pid = ctx.shared.get("prior_trial_balance_id")
        vault_st = (ctx.shared.get("prior_vault_snapshot") or {}).get("statements")
        if pid or vault_st:
            prior_pl = (
                _lines_for_statement(ctx.db, int(pid), ctx.tenant_id, IFRSStatementKind.profit_loss)
                if pid
                else _lines_from_vault_payload(vault_st, IFRSStatementKind.profit_loss)
            )
            pmap = {
                (getattr(x, "ifrs_section", "") or "", (getattr(x, "ifrs_line_item", "") or "").strip().casefold()): float(
                    getattr(x, "amount", 0) or 0
                )
                for x in (prior_pl or [])
            }
            pdf.cell(90, 5, f"Line (FY {cy_l})", border=0)
            pdf.cell(45, 5, "Current", border=0, ln=0)
            pdf.cell(45, 5, f"Prior FY {py_l}", border=0, ln=1)
            for li in pl_lines[:35]:
                sec = getattr(li, "ifrs_section", "") or ""
                nm = (li.ifrs_line_item or "").strip().casefold()
                pv = pmap.get((sec, nm), None)
                pdf.cell(90, 4, (li.ifrs_line_item or "")[:55], ln=0)
                pdf.cell(45, 4, f"{float(li.amount or 0):,.0f}", ln=0)
                pdf.cell(45, 4, f"{pv:,.0f}" if pv is not None else "-", ln=1)
        else:
            for li in pl_lines[:40]:
                pdf.cell(0, 5, f"{li.ifrs_line_item[:70]}: {float(li.amount or 0):,.2f}", ln=True)
        out_pdf = pdf.output(dest="S")
        if isinstance(out_pdf, str):
            out_pdf = out_pdf.encode("latin-1")
        pdf_path.write_bytes(out_pdf)

        rel = f"/api/ifrs/agentic/{ctx.public_run_id}/download"
        exports = {
            "xlsx_url": f"{rel}/xlsx",
            "docx_url": f"{rel}/docx",
            "pdf_url": f"{rel}/pdf",
        }
        ctx.shared["export_paths"] = {
            "xlsx": str(xlsx_path),
            "docx": str(docx_path),
            "pdf": str(pdf_path),
        }
        ctx.emit(self.agent_id, f"Packaged outputs to {out_dir}")
        return exports


def load_export_bytes(path: Path, kind: str) -> tuple[bytes, str, str]:
    if not path.exists():
        raise FileNotFoundError(kind)
    data = path.read_bytes()
    if kind == "xlsx":
        return data, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", path.name
    if kind == "pdf":
        return data, "application/pdf", path.name
    if kind == "docx":
        ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if path.suffix.lower() == ".txt":
            ct = "text/plain"
        return data, ct, path.name
    raise ValueError(kind)
