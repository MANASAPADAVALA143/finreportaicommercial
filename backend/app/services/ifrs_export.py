"""
IFRS Statement Export Service
==============================
Generates Excel (.xlsx), PDF (.pdf), and Word (.docx) exports of IFRS statements.

All functions return raw bytes — the caller (route handler) wraps them in a
StreamingResponse / FileResponse with the appropriate Content-Type.
"""
from __future__ import annotations

import io
import logging
from datetime import datetime
from decimal import Decimal
from typing import Any

from sqlalchemy.orm import Session

from app.models.ifrs_statement import (
    DisclosureNote,
    GeneratedStatement,
    IFRSStatementKind,
    StatementLineItem,
    TrialBalance,
)

logger = logging.getLogger(__name__)

# ── Helpers ────────────────────────────────────────────────────────────────────

STATEMENT_ORDER = [
    ("profit_loss", "Profit & Loss Statement"),
    ("financial_position", "Statement of Financial Position"),
    ("cash_flows", "Statement of Cash Flows"),
    ("equity", "Statement of Changes in Equity"),
    ("other_comprehensive_income", "Other Comprehensive Income"),
]


def _fmt(v: float | int | Decimal | None, negative_parens: bool = True) -> str:
    """Format a number with commas; wrap negatives in parentheses."""
    try:
        n = float(v or 0)
    except (TypeError, ValueError):
        return "-"
    if n < 0 and negative_parens:
        return f"({abs(n):,.2f})"
    return f"{n:,.2f}"


def _load_statements(trial_balance_id: int, db: Session) -> dict[str, list[StatementLineItem]]:
    """Return {statement_type: [sorted line items]}."""
    stmts = (
        db.query(GeneratedStatement)
        .filter(GeneratedStatement.trial_balance_id == trial_balance_id)
        .all()
    )
    result: dict[str, list[StatementLineItem]] = {}
    for s in stmts:
        items = sorted(s.line_items, key=lambda x: x.display_order)
        result[s.statement_type.value if hasattr(s.statement_type, "value") else s.statement_type] = items
    return result


def _load_notes(trial_balance_id: int, db: Session) -> list[DisclosureNote]:
    return (
        db.query(DisclosureNote)
        .filter(DisclosureNote.trial_balance_id == trial_balance_id)
        .order_by(DisclosureNote.note_number)
        .all()
    )


def _load_ct_bridge(trial_balance_id: int, db: Session) -> dict[str, Any] | None:
    try:
        from app.services.uae_ct_bridge import get_saved_ct_bridge
        return get_saved_ct_bridge(trial_balance_id, db)
    except Exception:
        return None


# ══════════════════════════════════════════════════════════════════════════════
# EXCEL EXPORT
# ══════════════════════════════════════════════════════════════════════════════

def export_to_excel(trial_balance_id: int, db: Session) -> bytes:
    """
    Generate a professional multi-sheet Excel workbook.
    Returns raw .xlsx bytes.
    """
    import openpyxl
    from openpyxl.styles import (
        Alignment,
        Border,
        Font,
        PatternFill,
        Side,
    )
    from openpyxl.utils import get_column_letter

    tb = db.query(TrialBalance).filter(TrialBalance.id == trial_balance_id).first()
    if not tb:
        raise ValueError(f"Trial balance {trial_balance_id} not found")

    company = tb.company_name
    period = str(tb.period_end) if tb.period_end else "N/A"
    currency = tb.currency or "AED"

    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # remove default sheet

    # ── Styles ────────────────────────────────────────────────────────────────
    def _header_font() -> Font:
        return Font(name="Calibri", bold=True, size=11, color="FFFFFF")

    def _title_font() -> Font:
        return Font(name="Calibri", bold=True, size=13, color="1F3864")

    def _subtotal_font() -> Font:
        return Font(name="Calibri", bold=True, size=10)

    def _normal_font() -> Font:
        return Font(name="Calibri", size=10)

    def _mono_font(bold: bool = False) -> Font:
        return Font(name="Courier New", size=10, bold=bold)

    NAVY = PatternFill("solid", fgColor="1F3864")
    GRAY = PatternFill("solid", fgColor="D9D9D9")
    LIGHT_BLUE = PatternFill("solid", fgColor="DCE6F1")
    thin_border = Border(
        bottom=Side(style="thin", color="BFBFBF"),
    )
    thick_border = Border(
        top=Side(style="medium", color="1F3864"),
        bottom=Side(style="medium", color="1F3864"),
    )

    def _setup_sheet(ws, title: str, stmt_type: str) -> None:
        ws.column_dimensions["A"].width = 48
        ws.column_dimensions["B"].width = 20
        ws.column_dimensions["C"].width = 16

        # Title block
        ws.merge_cells("A1:C1")
        ws["A1"] = company
        ws["A1"].font = _title_font()
        ws["A1"].alignment = Alignment(horizontal="left")

        ws.merge_cells("A2:C2")
        ws["A2"] = title
        ws["A2"].font = Font(name="Calibri", bold=True, size=11, color="44546A")

        ws.merge_cells("A3:C3")
        ws["A3"] = f"For the period ending {period} | Currency: {currency}"
        ws["A3"].font = Font(name="Calibri", size=9, italic=True, color="808080")

        # Column headers
        for col, label in [(1, "Line Item"), (2, f"Amount ({currency})"), (3, "Notes")]:
            cell = ws.cell(row=4, column=col, value=label)
            cell.font = _header_font()
            cell.fill = NAVY
            cell.alignment = Alignment(horizontal="center")

        ws.row_dimensions[1].height = 20
        ws.row_dimensions[4].height = 16

    statements = _load_statements(trial_balance_id, db)

    for stmt_key, stmt_label in STATEMENT_ORDER:
        line_items = statements.get(stmt_key)
        if not line_items:
            continue

        ws = wb.create_sheet(title=stmt_label[:31])  # Excel max 31 chars
        _setup_sheet(ws, stmt_label, stmt_key)

        row = 5
        current_section = None
        for li in line_items:
            if li.ifrs_section != current_section:
                current_section = li.ifrs_section
                # Section header row
                ws.merge_cells(f"A{row}:C{row}")
                cell = ws.cell(row=row, column=1, value=current_section.upper())
                cell.font = Font(name="Calibri", bold=True, size=9, color="44546A")
                cell.fill = GRAY
                cell.alignment = Alignment(horizontal="left", indent=1)
                ws.row_dimensions[row].height = 14
                row += 1

            indent = "  " * (li.indent_level or 0)
            label = f"{indent}{li.ifrs_line_item}"
            amount = float(li.amount or 0)
            is_bold = li.is_total or li.is_subtotal

            name_cell = ws.cell(row=row, column=1, value=label)
            amount_cell = ws.cell(row=row, column=2, value=amount)

            name_cell.font = _subtotal_font() if is_bold else _normal_font()
            amount_cell.font = _mono_font(bold=is_bold)
            amount_cell.number_format = '#,##0.00_);(#,##0.00)'
            amount_cell.alignment = Alignment(horizontal="right")
            if amount < 0:
                amount_cell.font = Font(
                    name="Courier New", size=10, bold=is_bold, color="C00000"
                )

            if li.is_total:
                for col in range(1, 4):
                    ws.cell(row=row, column=col).border = thick_border
                name_cell.fill = LIGHT_BLUE
                amount_cell.fill = LIGHT_BLUE
            elif li.is_subtotal:
                for col in range(1, 3):
                    ws.cell(row=row, column=col).border = thin_border

            row += 1

        ws.freeze_panes = "A5"
        ws.print_title_rows = "1:4"

    # ── UAE CT Bridge sheet ───────────────────────────────────────────────────
    ct = _load_ct_bridge(trial_balance_id, db)
    if ct:
        ws_ct = wb.create_sheet(title="UAE CT Bridge")
        ws_ct.column_dimensions["A"].width = 52
        ws_ct.column_dimensions["B"].width = 22

        ws_ct.merge_cells("A1:B1")
        ws_ct["A1"] = f"{company} — UAE Corporate Tax Bridge"
        ws_ct["A1"].font = _title_font()

        ws_ct.merge_cells("A2:B2")
        ws_ct["A2"] = f"Period: {ct.get('period_end', period)} | Prepared: {datetime.utcnow().strftime('%d %b %Y')}"
        ws_ct["A2"].font = Font(name="Calibri", size=9, italic=True, color="808080")

        for col, label in [(1, "Description"), (2, f"Amount ({ct.get('currency', 'AED')})")]:
            cell = ws_ct.cell(row=3, column=col, value=label)
            cell.font = _header_font()
            cell.fill = NAVY

        ct_rows = [
            ("IFRS Net Profit Before Tax", ct.get("ifrs_pbt", 0), False),
        ]
        for adj in ct.get("adjustments", []):
            sign = "Add: " if adj.get("add_back") else "Less: "
            ct_rows.append((f"  {sign}{adj['description']}", adj["amount"], False))
        ct_rows.append(("Taxable Income", ct.get("taxable_income", 0), True))
        ct_rows.append((
            f"UAE Corporate Tax @ {ct.get('ct_rate_pct', 9):.0f}%",
            ct.get("ct_liability", 0),
            True,
        ))

        row = 4
        for label, amount, bold in ct_rows:
            lc = ws_ct.cell(row=row, column=1, value=label)
            ac = ws_ct.cell(row=row, column=2, value=float(amount))
            lc.font = _subtotal_font() if bold else _normal_font()
            ac.font = _mono_font(bold=bold)
            ac.number_format = '#,##0.00_);(#,##0.00)'
            ac.alignment = Alignment(horizontal="right")
            if bold:
                for col in range(1, 3):
                    ws_ct.cell(row=row, column=col).border = thick_border
                lc.fill = LIGHT_BLUE
                ac.fill = LIGHT_BLUE
            row += 1

        notes_text = ct.get("rate_note", "")
        if ct.get("small_business_relief"):
            notes_text = "✅ Small Business Relief — 0% CT applies"
        elif ct.get("free_zone_eligible"):
            notes_text = "✅ Qualifying Free Zone Person — 0% CT on qualifying income"
        ws_ct.merge_cells(f"A{row+1}:B{row+1}")
        ws_ct.cell(row=row + 1, column=1, value=notes_text).font = Font(
            name="Calibri", size=9, italic=True, color="808080"
        )

    # ── Disclosure Notes sheet ────────────────────────────────────────────────
    notes = _load_notes(trial_balance_id, db)
    if notes:
        ws_notes = wb.create_sheet(title="Disclosure Notes")
        ws_notes.column_dimensions["A"].width = 18
        ws_notes.column_dimensions["B"].width = 100

        ws_notes.merge_cells("A1:B1")
        ws_notes["A1"] = f"{company} — Disclosure Notes"
        ws_notes["A1"].font = _title_font()

        row = 3
        for note in notes:
            content = note.user_edited_content or note.ai_generated_content or ""
            title_cell = ws_notes.cell(row=row, column=1, value=f"Note {note.note_number}")
            title_cell.font = Font(name="Calibri", bold=True, size=11, color="1F3864")
            ws_notes.cell(row=row, column=2, value=note.note_title).font = Font(
                name="Calibri", bold=True, size=11
            )
            row += 1
            ws_notes.merge_cells(f"A{row}:B{row}")
            content_cell = ws_notes.cell(row=row, column=1, value=content)
            content_cell.font = Font(name="Calibri", size=9)
            content_cell.alignment = Alignment(wrap_text=True)
            ws_notes.row_dimensions[row].height = max(60, len(content) // 6)
            row += 2

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# PDF EXPORT
# ══════════════════════════════════════════════════════════════════════════════

def export_to_pdf(trial_balance_id: int, db: Session) -> bytes:
    """
    Generate a professional multi-page PDF using ReportLab.
    Returns raw .pdf bytes.
    """
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        HRFlowable,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    tb = db.query(TrialBalance).filter(TrialBalance.id == trial_balance_id).first()
    if not tb:
        raise ValueError(f"Trial balance {trial_balance_id} not found")

    company = tb.company_name
    period = str(tb.period_end) if tb.period_end else "N/A"
    currency = tb.currency or "AED"

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2.2 * cm,
        rightMargin=2.2 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        title=f"{company} — IFRS Financial Statements",
        author="FinReportAI",
    )

    styles = getSampleStyleSheet()
    NAVY = colors.HexColor("#1F3864")
    GRAY = colors.HexColor("#F3F4F6")
    DARK_GRAY = colors.HexColor("#44546A")
    RED = colors.HexColor("#C00000")

    s_company = ParagraphStyle("Company", fontName="Helvetica-Bold", fontSize=16, textColor=NAVY, spaceAfter=4)
    s_subtitle = ParagraphStyle("Subtitle", fontName="Helvetica", fontSize=10, textColor=DARK_GRAY, spaceAfter=2)
    s_stmt_title = ParagraphStyle("StmtTitle", fontName="Helvetica-Bold", fontSize=13, textColor=NAVY, spaceBefore=12, spaceAfter=6)
    s_section = ParagraphStyle("Section", fontName="Helvetica-Bold", fontSize=9, textColor=DARK_GRAY)
    s_note_title = ParagraphStyle("NoteTitle", fontName="Helvetica-Bold", fontSize=11, textColor=NAVY, spaceBefore=10, spaceAfter=4)
    s_note_body = ParagraphStyle("NoteBody", fontName="Helvetica", fontSize=9, leading=13, spaceAfter=6)
    s_footer = ParagraphStyle("Footer", fontName="Helvetica-Oblique", fontSize=7, textColor=colors.gray, alignment=TA_CENTER)

    story: list = []

    def _page_header() -> list:
        return [
            Paragraph(company, s_company),
            Paragraph(f"IFRS Financial Statements — Period ending {period}", s_subtitle),
            Paragraph(f"Currency: {currency} | Generated: {datetime.utcnow().strftime('%d %b %Y at %H:%M UTC')}", s_subtitle),
            HRFlowable(width="100%", thickness=2, color=NAVY, spaceAfter=8),
        ]

    story.extend(_page_header())

    statements = _load_statements(trial_balance_id, db)

    for stmt_key, stmt_label in STATEMENT_ORDER:
        line_items = statements.get(stmt_key)
        if not line_items:
            continue

        story.append(Paragraph(stmt_label.upper(), s_stmt_title))

        tdata = [["Line Item", f"Amount ({currency})"]]
        tstyle = [
            ("BACKGROUND", (0, 0), (-1, 0), NAVY),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 9),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F9FAFB")]),
            ("FONTSIZE", (0, 1), (-1, -1), 8),
            ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
            ("ALIGN", (1, 0), (1, -1), "RIGHT"),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING", (0, 0), (0, -1), 6),
            ("GRID", (0, 0), (-1, 0), 0.5, NAVY),
            ("LINEBELOW", (0, 0), (-1, 0), 1, NAVY),
        ]

        current_section = None
        row_idx = 1

        for li in line_items:
            if li.ifrs_section != current_section:
                current_section = li.ifrs_section
                tdata.append([Paragraph(f"<b>{current_section.upper()}</b>", s_section), ""])
                tstyle.append(("BACKGROUND", (0, row_idx), (-1, row_idx), GRAY))
                tstyle.append(("SPAN", (0, row_idx), (1, row_idx)))
                row_idx += 1

            indent = " " * (4 * (li.indent_level or 0))
            label = f"{indent}{li.ifrs_line_item}"
            amount = float(li.amount or 0)
            amount_str = _fmt(amount)

            if li.is_total:
                tdata.append([
                    Paragraph(f"<b>{label}</b>", ParagraphStyle("TotL", fontName="Helvetica-Bold", fontSize=9)),
                    Paragraph(f"<b>{amount_str}</b>", ParagraphStyle("TotR", fontName="Helvetica-Bold", fontSize=9, alignment=TA_RIGHT)),
                ])
                tstyle.append(("LINEABOVE", (0, row_idx), (-1, row_idx), 1.5, NAVY))
                tstyle.append(("LINEBELOW", (0, row_idx), (-1, row_idx), 1.5, NAVY))
                tstyle.append(("BACKGROUND", (0, row_idx), (-1, row_idx), colors.HexColor("#DCE6F1")))
            elif li.is_subtotal:
                tdata.append([
                    Paragraph(f"<b>{label}</b>", ParagraphStyle("SubL", fontName="Helvetica-Bold", fontSize=8)),
                    Paragraph(f"<b>{amount_str}</b>", ParagraphStyle("SubR", fontName="Helvetica-Bold", fontSize=8, alignment=TA_RIGHT)),
                ])
                tstyle.append(("LINEABOVE", (0, row_idx), (-1, row_idx), 0.5, DARK_GRAY))
            else:
                color_amount = RED if amount < 0 else colors.black
                tdata.append([label, Paragraph(f'<font color="#{color_amount.hexval()[2:]}">{amount_str}</font>',
                                               ParagraphStyle("Amt", fontName="Courier", fontSize=8, alignment=TA_RIGHT))])

            row_idx += 1

        col_widths = [12 * cm, 4 * cm]
        t = Table(tdata, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle(tstyle))
        story.append(t)
        story.append(PageBreak())

    # ── UAE CT Bridge ─────────────────────────────────────────────────────────
    ct = _load_ct_bridge(trial_balance_id, db)
    if ct:
        story.append(Paragraph("UAE CORPORATE TAX BRIDGE", s_stmt_title))
        ct_rows = [["Description", f"Amount ({ct.get('currency', 'AED')})"]]
        ct_rows.append(["IFRS Net Profit Before Tax", _fmt(ct.get("ifrs_pbt", 0))])
        for adj in ct.get("adjustments", []):
            sign = "Add: " if adj.get("add_back") else "Less: "
            ct_rows.append([f"   {sign}{adj['description']}", _fmt(adj.get("amount", 0))])
        ct_rows.append(["Taxable Income", _fmt(ct.get("taxable_income", 0))])
        ct_rows.append([f"UAE Corporate Tax @ {ct.get('ct_rate_pct', 9):.0f}%", _fmt(ct.get("ct_liability", 0))])

        ct_style = [
            ("BACKGROUND", (0, 0), (-1, 0), NAVY),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("ALIGN", (1, 0), (1, -1), "RIGHT"),
            ("LINEBELOW", (0, -1), (-1, -1), 2, NAVY),
            ("LINEABOVE", (0, -1), (-1, -1), 1.5, NAVY),
            ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#DCE6F1")),
            ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
            ("BACKGROUND", (0, -2), (-1, -2), colors.HexColor("#DCE6F1")),
            ("FONTNAME", (0, -2), (-1, -2), "Helvetica-Bold"),
        ]
        ct_t = Table(ct_rows, colWidths=[12 * cm, 4 * cm])
        ct_t.setStyle(TableStyle(ct_style))
        story.append(ct_t)
        story.append(Spacer(1, 0.5 * cm))
        story.append(Paragraph(ct.get("rate_note", ""), s_note_body))
        story.append(PageBreak())

    # ── Disclosure Notes ──────────────────────────────────────────────────────
    notes = _load_notes(trial_balance_id, db)
    if notes:
        story.append(Paragraph("NOTES TO THE FINANCIAL STATEMENTS", s_stmt_title))
        for note in notes:
            content = note.user_edited_content or note.ai_generated_content or "(not yet generated)"
            story.append(Paragraph(f"Note {note.note_number}: {note.note_title}", s_note_title))
            story.append(Paragraph(content.replace("\n", "<br/>"), s_note_body))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey, spaceAfter=6))

    # ── Page footer via canvas callback ──────────────────────────────────────
    def _add_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica-Oblique", 7)
        canvas.setFillColor(colors.gray)
        footer_text = f"{company} | IFRS Financial Statements | Page {doc.page} | Confidential"
        canvas.drawCentredString(A4[0] / 2, 1.5 * cm, footer_text)
        canvas.restoreState()

    doc.build(story, onFirstPage=_add_footer, onLaterPages=_add_footer)
    return buffer.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# WORD EXPORT
# ══════════════════════════════════════════════════════════════════════════════

def export_to_word(trial_balance_id: int, db: Session) -> bytes:
    """
    Generate a professional Word document using python-docx.
    Returns raw .docx bytes.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    import lxml.etree as etree

    tb = db.query(TrialBalance).filter(TrialBalance.id == trial_balance_id).first()
    if not tb:
        raise ValueError(f"Trial balance {trial_balance_id} not found")

    company = tb.company_name
    period = str(tb.period_end) if tb.period_end else "N/A"
    currency = tb.currency or "AED"

    NAVY_RGB = RGBColor(0x1F, 0x38, 0x64)
    GRAY_RGB = RGBColor(0xD9, 0xD9, 0xD9)

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    def _add_cover() -> None:
        p = doc.add_heading(company, level=1)
        p.runs[0].font.color.rgb = NAVY_RGB
        doc.add_paragraph(f"IFRS Financial Statements")
        doc.add_paragraph(f"For the period ending: {period}")
        doc.add_paragraph(f"Currency: {currency}")
        doc.add_paragraph(f"Prepared by: FinReportAI | {datetime.utcnow().strftime('%d %B %Y')}")
        doc.add_page_break()

    def _add_statement(stmt_label: str, line_items: list[StatementLineItem]) -> None:
        h = doc.add_heading(stmt_label, level=2)
        h.runs[0].font.color.rgb = NAVY_RGB

        doc.add_paragraph(f"{company} | Period ending {period} | {currency}")

        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.columns[0].width = Inches(4.5)
        tbl.columns[1].width = Inches(1.8)

        # Header row
        hdr = tbl.rows[0].cells
        hdr[0].text = "Line Item"
        hdr[1].text = f"Amount ({currency})"
        for cell in hdr:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_bg(cell, "1F3864")
        hdr[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        current_section = None
        for li in line_items:
            if li.ifrs_section != current_section:
                current_section = li.ifrs_section
                sec_row = tbl.add_row()
                sec_row.cells[0].merge(sec_row.cells[1])
                sec_row.cells[0].text = current_section.upper()
                r = sec_row.cells[0].paragraphs[0].runs[0]
                r.font.bold = True
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(0x44, 0x54, 0x6A)
                _set_cell_bg(sec_row.cells[0], "D9D9D9")

            row_cells = tbl.add_row().cells
            indent = "  " * (li.indent_level or 0)
            amount = float(li.amount or 0)
            is_bold = li.is_total or li.is_subtotal

            row_cells[0].text = f"{indent}{li.ifrs_line_item}"
            row_cells[1].text = _fmt(amount)
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            for cell in row_cells:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = is_bold
                    run.font.size = Pt(9)
                    if amount < 0:
                        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

            if li.is_total:
                _set_cell_bg(row_cells[0], "DCE6F1")
                _set_cell_bg(row_cells[1], "DCE6F1")

        doc.add_page_break()

    def _set_cell_bg(cell, hex_color: str) -> None:
        """Set cell background color via XML."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = etree.SubElement(tcPr, qn("w:shd"))
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)

    _add_cover()

    statements = _load_statements(trial_balance_id, db)
    for stmt_key, stmt_label in STATEMENT_ORDER:
        line_items = statements.get(stmt_key)
        if line_items:
            _add_statement(stmt_label, line_items)

    # ── UAE CT Bridge ─────────────────────────────────────────────────────────
    ct = _load_ct_bridge(trial_balance_id, db)
    if ct:
        h = doc.add_heading("UAE Corporate Tax Bridge", level=2)
        h.runs[0].font.color.rgb = NAVY_RGB

        ct_tbl = doc.add_table(rows=1, cols=2)
        ct_tbl.style = "Table Grid"
        ct_tbl.columns[0].width = Inches(4.5)
        ct_tbl.columns[1].width = Inches(1.8)

        hdr = ct_tbl.rows[0].cells
        hdr[0].text = "Description"
        hdr[1].text = f"Amount ({ct.get('currency', 'AED')})"
        for cell in hdr:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_bg(cell, "1F3864")

        def _ct_row(label: str, amount: float, bold: bool = False) -> None:
            r = ct_tbl.add_row().cells
            r[0].text = label
            r[1].text = _fmt(amount)
            r[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for cell in r:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = bold
                    run.font.size = Pt(9)
            if bold:
                _set_cell_bg(r[0], "DCE6F1")
                _set_cell_bg(r[1], "DCE6F1")

        _ct_row("IFRS Net Profit Before Tax", ct.get("ifrs_pbt", 0))
        for adj in ct.get("adjustments", []):
            sign = "Add: " if adj.get("add_back") else "Less: "
            _ct_row(f"   {sign}{adj['description']}", adj.get("amount", 0))
        _ct_row("Taxable Income", ct.get("taxable_income", 0), bold=True)
        _ct_row(
            f"UAE Corporate Tax @ {ct.get('ct_rate_pct', 9):.0f}%",
            ct.get("ct_liability", 0),
            bold=True,
        )

        doc.add_paragraph()
        note_p = doc.add_paragraph(ct.get("rate_note", ""))
        note_p.runs[0].font.italic = True
        note_p.runs[0].font.size = Pt(9)
        doc.add_page_break()

    # ── Disclosure Notes ──────────────────────────────────────────────────────
    notes = _load_notes(trial_balance_id, db)
    if notes:
        h = doc.add_heading("Notes to the Financial Statements", level=2)
        h.runs[0].font.color.rgb = NAVY_RGB
        for note in notes:
            content = note.user_edited_content or note.ai_generated_content or "(not yet generated)"
            nh = doc.add_heading(f"Note {note.note_number}: {note.note_title}", level=3)
            nh.runs[0].font.color.rgb = NAVY_RGB
            doc.add_paragraph(content)
            doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
