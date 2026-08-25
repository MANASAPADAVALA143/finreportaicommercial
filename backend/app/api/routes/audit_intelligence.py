"""Audit Intelligence — five Claude-powered audit agents."""

from __future__ import annotations

import io
import json
from datetime import datetime, timezone
from typing import Any

import pandas as pd
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import Response
from openpyxl import Workbook
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session

from app.core.database import get_db
from app.models.audit_intelligence import AuditRun
from app.models.client_data import ApInvoice, CtReturn, GulftaxTransaction
from app.models.uae_accounting_full import (
    UAEAccount,
    UAEJournalEntry,
    UAEJournalLine,
    UAESalesInvoice,
)
from app.services.audit_claude import invoke_audit_json
from app.services.audit_intelligence_helpers import (
    read_upload_as_csv_text,
    read_upload_as_dataframe,
    result_summary_for_agent,
)
from app.services.audit_pdf_report import build_audit_pdf_bytes
from app.services.audit_command_center_service import parse_period_dates
from app.services.llm_service import invoke
from app.services.r2r_pattern_engine import R2RPatternEngine
from docx import Document

router = APIRouter(prefix="/api/audit", tags=["audit-intelligence"])


GOING_CONCERN_SYSTEM_PROMPT = (
    "You are a senior auditor assessing going concern under ISA 570 and IAS 1 paragraph 25. "
    "Based on the indicators provided, write a concise going concern assessment paragraph "
    "(3-5 sentences) in professional audit language. If material uncertainty exists, state it "
    "explicitly. Do not use bullet points."
)

# --- System prompts (JSON-only responses) ---

SYS_EVIDENCE = """You are a Big 4 senior auditor with 15 years experience. Analyse this transaction data applying ISA (International Standards on Auditing) guidelines.
Identify high-risk transactions, unusual patterns, round number transactions, period-end entries, and related party transactions. Apply risk-based sampling.
Format output as structured JSON only with this exact shape:
{
  "summary": {
    "total_transactions": number,
    "selected_for_testing": number,
    "high_risk_count": number,
    "total_value_selected": number
  },
  "evidence_checklist": [
    {
      "transaction_ref": string,
      "account": string,
      "amount": number,
      "risk_level": "High" | "Medium" | "Low",
      "evidence_required": string,
      "status": "Pending"
    }
  ],
  "audit_findings": string,
  "recommendations": string
}"""

SYS_EVIDENCE_R2R = """You are a Big 4 senior auditor with 15 years experience. You have received Journal Entry pattern analysis results from an R2R engine.

Your job is to:
1. Review all flagged patterns
2. Assess audit risk for each finding
3. Determine which entries need audit evidence
4. Identify if patterns suggest fraud risk
5. Generate ISA 530 compliant sample selection
6. Write audit findings in professional language

Apply ISA (International Standards on Auditing):
- ISA 240: Fraud risk assessment
- ISA 315: Risk identification
- ISA 530: Audit sampling
- ISA 240: Journal entry testing

Output structured JSON only with this exact shape:
{
  "audit_summary": {
    "total_entries_analysed": number,
    "patterns_detected": number,
    "high_risk_entries": number,
    "medium_risk_entries": number,
    "fraud_indicators": number,
    "audit_risk_rating": "Low" | "Medium" | "High" | "Critical"
  },
  "r2r_patterns_found": [
    {
      "pattern_type": string,
      "entries_affected": number,
      "total_value": number,
      "risk_level": "High" | "Medium" | "Low",
      "description": string,
      "entries": [
        {
          "date": string,
          "account": string,
          "amount": number,
          "description": string,
          "reference": string,
          "flag_reason": string
        }
      ]
    }
  ],
  "audit_evidence_required": [
    {
      "priority": number,
      "entry_reference": string,
      "amount": number,
      "risk_level": string,
      "pattern_detected": string,
      "evidence_needed": string,
      "isa_reference": string,
      "status": "Pending"
    }
  ],
  "fraud_risk_assessment": {
    "fraud_risk_level": string,
    "indicators_found": [ string ],
    "recommended_procedures": [ string ]
  },
  "auditor_findings": string,
  "management_letter_points": [ string ]
}"""

SYS_IFRS = """You are an IFRS technical expert. Review the provided financial information against the specified IFRS standard.
Check every disclosure requirement, measurement basis, and presentation requirement. Be precise and cite specific paragraphs of the standard.
Output structured JSON only with this exact shape:
{
  "compliance_score": number,
  "standard_checked": string,
  "compliant_items": [ string ],
  "non_compliant_items": [
    {
      "area": string,
      "requirement": string,
      "gap_found": string,
      "severity": "Critical" | "Major" | "Minor",
      "recommendation": string
    }
  ],
  "disclosure_gaps": [ string ],
  "overall_opinion": string
}"""

SYS_CONTROLS = """You are an internal audit director. Assess the described process for internal control weaknesses.
Apply COSO framework. Identify segregation of duties issues, authorisation gaps, and reconciliation weaknesses. Rate risks and provide actionable recommendations.
Output structured JSON only with this exact shape:
{
  "controls_identified": [
    {
      "control_name": string,
      "control_type": "Preventive" | "Detective" | "Corrective",
      "effectiveness": "Strong" | "Adequate" | "Weak",
      "test_result": "Pass" | "Fail"
    }
  ],
  "gaps_identified": [
    {
      "gap": string,
      "risk_level": "High" | "Medium" | "Low",
      "financial_impact": string,
      "recommendation": string
    }
  ],
  "overall_control_rating": string,
  "priority_actions": [ string ]
}"""

SYS_SOX = """You are a SOX compliance expert. Evaluate these control test results against SOX Section 404 requirements.
Classify all deficiencies per PCAOB standards. Determine if material weaknesses exist. Draft management conclusion language.
Output structured JSON only with this exact shape:
{
  "sox_opinion": "Effective" | "Qualified" | "Adverse",
  "material_weaknesses": [ string ],
  "significant_deficiencies": [ string ],
  "control_deficiencies": [ string ],
  "management_conclusion": string,
  "remediation_plan": [ string ]
}"""

SYS_AML = """You are an AML compliance specialist. Analyse these transactions for money laundering patterns including structuring, layering, smurfing,
round-trip transactions, and unusual velocity. Apply FATF guidelines. Generate risk scores. Flag SAR-worthy transactions with justification.
Output structured JSON only with this exact shape:
{
  "summary": {
    "total_scanned": number,
    "flagged_count": number,
    "sar_required": number,
    "total_flagged_value": number
  },
  "flagged_transactions": [
    {
      "reference": string,
      "amount": number,
      "risk_score": number,
      "red_flags": [ string ],
      "pattern_detected": string,
      "sar_recommended": boolean
    }
  ],
  "overall_risk_rating": "Low" | "Medium" | "High" | "Critical",
  "regulatory_obligations": string
}"""


def _pick_col(df: pd.DataFrame, names: list[str]) -> str | None:
    lookup = {c.lower().strip(): c for c in df.columns}
    for name in names:
        if name in lookup:
            return lookup[name]
    return None


def _num(val: Any) -> float:
    try:
        if val is None:
            return 0.0
        return float(val)
    except (TypeError, ValueError):
        return 0.0


def _normalize_je_df(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    date_col = _pick_col(out, ["date", "posting_date", "posted_date", "voucher_date"])
    acct_col = _pick_col(out, ["account", "gl_account", "ledger", "account_name"])
    debit_col = _pick_col(out, ["debit", "dr"])
    credit_col = _pick_col(out, ["credit", "cr"])
    amount_col = _pick_col(out, ["amount", "amt", "value"])
    desc_col = _pick_col(out, ["description", "narration", "memo", "details"])
    ref_col = _pick_col(out, ["reference", "ref", "voucher_no", "je_id", "entry_id"])
    user_col = _pick_col(out, ["user", "posted_by", "preparer", "created_by"])
    time_col = _pick_col(out, ["time", "posting_time", "posted_time"])
    approval_col = _pick_col(out, ["approval_reference", "approval_ref", "approval", "approved_by"])
    auth_limit_col = _pick_col(out, ["authorisation_limit", "authorization_limit", "auth_limit"])
    entity_col = _pick_col(out, ["entity", "company", "legal_entity"])
    interco_col = _pick_col(out, ["intercompany_ref", "interco_ref", "elimination_ref"])

    out["_date"] = pd.to_datetime(out[date_col], errors="coerce") if date_col else pd.NaT
    out["_account"] = out[acct_col].astype(str) if acct_col else ""
    out["_description"] = out[desc_col].fillna("").astype(str) if desc_col else ""
    out["_reference"] = out[ref_col].fillna("").astype(str) if ref_col else [f"JE-{i+1:04d}" for i in range(len(out))]
    out["_user"] = out[user_col].fillna("").astype(str) if user_col else ""
    out["_approval"] = out[approval_col].fillna("").astype(str) if approval_col else ""
    out["_entity"] = out[entity_col].fillna("").astype(str) if entity_col else ""
    out["_interco"] = out[interco_col].fillna("").astype(str) if interco_col else ""

    d = pd.to_numeric(out[debit_col], errors="coerce").fillna(0).abs() if debit_col else pd.Series(0, index=out.index)
    c = pd.to_numeric(out[credit_col], errors="coerce").fillna(0).abs() if credit_col else pd.Series(0, index=out.index)
    if amount_col:
        a = pd.to_numeric(out[amount_col], errors="coerce").fillna(0).abs()
        out["_amount"] = a.where(a > 0, d.where(d > 0, c))
    else:
        out["_amount"] = d.where(d > 0, c)
    out["_debit"] = d
    out["_credit"] = c

    if time_col:
        parsed_time = pd.to_datetime(out[time_col], errors="coerce")
        out["_hour"] = parsed_time.dt.hour.fillna(12).astype(int)
        out["_minute"] = parsed_time.dt.minute.fillna(0).astype(int)
    else:
        out["_hour"] = out["_date"].dt.hour.fillna(12).astype(int)
        out["_minute"] = out["_date"].dt.minute.fillna(0).astype(int)

    if auth_limit_col:
        out["_authorisation_limit"] = pd.to_numeric(out[auth_limit_col], errors="coerce").fillna(0)
    else:
        out["_authorisation_limit"] = 0.0
    return out


def _entry_obj(row: pd.Series, reason: str) -> dict[str, Any]:
    return {
        "date": row.get("_date").strftime("%Y-%m-%d") if pd.notna(row.get("_date")) else "",
        "account": str(row.get("_account") or ""),
        "amount": float(row.get("_amount") or 0),
        "description": str(row.get("_description") or ""),
        "reference": str(row.get("_reference") or ""),
        "flag_reason": reason,
    }


def _detect_audit_patterns(df: pd.DataFrame, authorisation_limit: float, period_end: datetime) -> list[dict[str, Any]]:
    patterns: list[dict[str, Any]] = []
    norm = _normalize_je_df(df)

    def add_pattern(name: str, risk: str, desc: str, rows: pd.DataFrame, reason: str) -> None:
        if rows.empty:
            return
        entries = [_entry_obj(r, reason) for _, r in rows.iterrows()]
        patterns.append(
            {
                "pattern_type": name,
                "entries_affected": len(entries),
                "total_value": round(float(rows["_amount"].sum()), 2),
                "risk_level": risk,
                "description": desc,
                "entries": entries[:50],
            }
        )

    dup_mask = norm.duplicated(subset=["_date", "_account", "_amount"], keep=False) & (norm["_amount"] > 0)
    add_pattern(
        "Duplicate Entries",
        "High",
        "Same account, date, and amount posted multiple times.",
        norm[dup_mask],
        "Same account + date + amount appears more than once.",
    )

    round_mask = (norm["_amount"] > 0) & ((norm["_amount"] % 1000 == 0) | (norm["_amount"] % 10000 == 0))
    add_pattern(
        "Round Number Transactions",
        "High",
        "Round-number postings ending in 000/0000.",
        norm[round_mask],
        "Amount ends in 000 or 0000.",
    )

    period_mask = norm["_date"].notna() & (
        (norm["_date"].dt.day >= 28)
        | ((period_end is not None) & (norm["_date"] > pd.Timestamp(period_end)))
    )
    add_pattern(
        "Period End Entries",
        "High",
        "Entries posted near period close or after period end.",
        norm[period_mask & (norm["_amount"] > 0)],
        "Posted on/after period-end window (28th-31st or after close).",
    )

    policy_mask = (norm["_hour"] < 8) | (norm["_hour"] > 20) | (norm["_approval"].str.strip() == "")
    if authorisation_limit > 0:
        policy_mask = policy_mask | (norm["_amount"] > authorisation_limit)
    add_pattern(
        "Policy Breach Detection",
        "High",
        "Outside hours, missing approval, or over authorization threshold.",
        norm[policy_mask & (norm["_amount"] > 0)],
        "Potential policy breach (hours/approval/limit).",
    )

    acct = norm["_account"].str.lower()
    unusual_mask = ((acct.str.contains("revenue")) & (norm["_debit"] > 0)) | (
        (acct.str.contains("expense")) & (norm["_credit"] > 0)
    )
    add_pattern(
        "Unusual Account Combinations",
        "Medium",
        "Debit to revenue or credit to expense patterns.",
        norm[unusual_mask & (norm["_amount"] > 0)],
        "Unusual debit/credit direction for account class.",
    )

    rapid_idx: list[int] = []
    if (norm["_user"].str.strip() != "").any() and norm["_date"].notna().any():
        tmp = norm.copy()
        tmp["_dt"] = pd.to_datetime(tmp["_date"].dt.strftime("%Y-%m-%d") + " " + tmp["_hour"].astype(str) + ":" + tmp["_minute"].astype(str), errors="coerce")
        for _, group in tmp.groupby("_user"):
            group = group.sort_values("_dt")
            if len(group) < 6:
                continue
            times = group["_dt"].tolist()
            idxs = group.index.tolist()
            for i in range(len(times)):
                if pd.isna(times[i]):
                    continue
                c = 1
                for j in range(i + 1, len(times)):
                    if pd.isna(times[j]):
                        continue
                    if (times[j] - times[i]).total_seconds() <= 600:
                        c += 1
                    else:
                        break
                if c > 5:
                    rapid_idx.extend(idxs[i : i + c])
    rapid_mask = norm.index.isin(rapid_idx)
    add_pattern(
        "Rapid Succession Entries",
        "High",
        "More than 5 entries by same user in 10 minutes.",
        norm[rapid_mask & (norm["_amount"] > 0)],
        "User velocity exceeds 5 entries within 10 minutes.",
    )

    interco_mask = norm["_account"].str.contains("interco|intercompany", case=False, na=False)
    interco_rows = norm[interco_mask & (norm["_amount"] > 0)].copy()
    if not interco_rows.empty:
        imbalanced = []
        for _, row in interco_rows.iterrows():
            ref = str(row.get("_interco") or row.get("_reference") or "")
            if not ref:
                imbalanced.append(row.name)
                continue
            matched = interco_rows[(interco_rows["_reference"] == ref) | (interco_rows["_interco"] == ref)]
            if len(matched) < 2:
                imbalanced.append(row.name)
        add_pattern(
            "Intercompany Imbalances",
            "High",
            "Intercompany postings without clear balancing match.",
            norm.loc[imbalanced],
            "Intercompany entry missing matching elimination/balance pair.",
        )

    missing_desc_mask = (norm["_description"].str.strip() == "") | (norm["_reference"].str.strip() == "")
    add_pattern(
        "Missing Descriptions",
        "Medium",
        "Blank narration or missing reference details.",
        norm[missing_desc_mask & (norm["_amount"] > 0)],
        "Description/reference field is blank.",
    )
    return patterns


def _run_r2r_engine(df: pd.DataFrame, materiality_amount: float) -> dict[str, Any]:
    engine = R2RPatternEngine()
    r2r_result = engine.analyse(df, sensitivity="balanced", materiality_amount=materiality_amount, materiality_pct=0)
    if isinstance(r2r_result, dict) and r2r_result.get("error"):
        raise HTTPException(status_code=400, detail=str(r2r_result["error"]))
    return r2r_result


def _merge_evidence_result(
    base_result: dict[str, Any],
    *,
    patterns: list[dict[str, Any]],
    total_entries: int,
    high_risk_entries: int,
    medium_risk_entries: int,
) -> dict[str, Any]:
    output = base_result if isinstance(base_result, dict) else {}
    output.setdefault("r2r_patterns_found", patterns)
    output.setdefault("audit_evidence_required", [])
    output.setdefault("fraud_risk_assessment", {"fraud_risk_level": "Medium", "indicators_found": [], "recommended_procedures": []})
    output.setdefault("auditor_findings", "")
    output.setdefault("management_letter_points", [])
    fraud_indicators = len(output.get("fraud_risk_assessment", {}).get("indicators_found", []))
    risk_rating = output.get("audit_summary", {}).get("audit_risk_rating")
    if not risk_rating:
        risk_rating = "Critical" if high_risk_entries >= 10 else ("High" if high_risk_entries >= 4 else ("Medium" if medium_risk_entries >= 3 else "Low"))
    output["audit_summary"] = {
        "total_entries_analysed": int(total_entries),
        "patterns_detected": int(len(patterns)),
        "high_risk_entries": int(high_risk_entries),
        "medium_risk_entries": int(medium_risk_entries),
        "fraud_indicators": int(fraud_indicators),
        "audit_risk_rating": str(risk_rating),
    }
    return output


def _persist(
    db: Session,
    *,
    agent_type: str,
    client_name: str | None,
    file_name: str | None,
    result: dict[str, Any],
) -> AuditRun:
    row = AuditRun(
        agent_type=agent_type,
        client_name=(client_name or None),
        file_name=file_name,
        run_timestamp=datetime.now(timezone.utc),
        result_summary=result_summary_for_agent(agent_type, result),
        full_result=result,
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    return row


@router.post("/evidence-collector")
async def audit_evidence_collector(
    file: UploadFile = File(...),
    audit_period_start: str = Form(...),
    audit_period_end: str = Form(...),
    risk_threshold: str = Form("Medium"),
    materiality_threshold: float = Form(0),
    authorisation_limit: float = Form(0),
    client_name: str | None = Form(None),
    db: Session = Depends(get_db),
):
    df, row_count = await read_upload_as_dataframe(file)
    csv_text = df.to_csv(index=False)
    period_end_dt = datetime.fromisoformat(audit_period_end)

    r2r_raw = _run_r2r_engine(df, materiality_amount=float(materiality_threshold or 0))
    scored = r2r_raw.get("entries_scored", []) if isinstance(r2r_raw, dict) else []
    high_risk_entries = sum(1 for x in scored if str(x.get("risk_level", "")).upper() == "HIGH")
    medium_risk_entries = sum(1 for x in scored if str(x.get("risk_level", "")).upper() == "MEDIUM")
    patterns = _detect_audit_patterns(df, float(authorisation_limit or 0), period_end_dt)

    user = f"""Audit period: {audit_period_start} to {audit_period_end}
Risk threshold (sampling focus): {risk_threshold}
Materiality threshold amount: {materiality_threshold}
Authorisation limit amount: {authorisation_limit}
Uploaded rows: {row_count}

R2R engine summary:
{json.dumps(r2r_raw.get("summary", {}), default=str)}

R2R pattern findings:
{json.dumps(patterns, default=str)}

Raw data (CSV):
{csv_text}
"""
    base_result = invoke_audit_json(SYS_EVIDENCE_R2R, user)
    result = _merge_evidence_result(
        base_result,
        patterns=patterns,
        total_entries=row_count,
        high_risk_entries=high_risk_entries,
        medium_risk_entries=medium_risk_entries,
    )
    run = _persist(
        db,
        agent_type="evidence-collector",
        client_name=client_name,
        file_name=file.filename,
        result=result,
    )
    return {"run_id": run.id, "result": result}


@router.post("/ifrs-checker")
async def audit_ifrs_checker(
    ifrs_standard: str = Form(...),
    entity_type: str = Form(...),
    client_name: str | None = Form(None),
    financial_text: str = Form(""),
    file: UploadFile | None = File(None),
    db: Session = Depends(get_db),
):
    file_name = None
    parts = [
        f"IFRS standard to check: {ifrs_standard}",
        f"Entity type: {entity_type}",
    ]
    if financial_text.strip():
        parts.append(f"Financial information (text):\n{financial_text.strip()}")
    if file and file.filename:
        csv_text, row_count = await read_upload_as_csv_text(file)
        parts.append(f"Uploaded data ({file.filename}, {row_count} rows):\n{csv_text}")
        file_name = file.filename
    if len(parts) < 3:
        raise HTTPException(status_code=400, detail="Provide financial_text and/or a CSV/XLSX file.")
    user = "\n\n".join(parts)
    result = invoke_audit_json(SYS_IFRS, user)
    run = _persist(
        db,
        agent_type="ifrs-checker",
        client_name=client_name,
        file_name=file_name,
        result=result,
    )
    return {"run_id": run.id, "result": result}


@router.post("/controls-tester")
async def audit_controls_tester(
    process_description: str = Form(...),
    control_type: str = Form(...),
    company_size: str = Form(...),
    client_name: str | None = Form(None),
    db: Session = Depends(get_db),
):
    user = f"""Control category: {control_type}
Company size: {company_size}

Process description:
{process_description}
"""
    result = invoke_audit_json(SYS_CONTROLS, user)
    run = _persist(
        db,
        agent_type="controls-tester",
        client_name=client_name,
        file_name=None,
        result=result,
    )
    return {"run_id": run.id, "result": result}


@router.post("/sox-checker")
async def audit_sox_checker(
    file: UploadFile = File(...),
    section: str = Form(...),
    quarter: str = Form(...),
    client_name: str | None = Form(None),
    db: Session = Depends(get_db),
):
    csv_text, row_count = await read_upload_as_csv_text(file)
    user = f"""SOX scope: {section}
Quarter: {quarter}
Control testing results ({row_count} rows):
{csv_text}
"""
    result = invoke_audit_json(SYS_SOX, user)
    run = _persist(
        db,
        agent_type="sox-checker",
        client_name=client_name,
        file_name=file.filename,
        result=result,
    )
    return {"run_id": run.id, "result": result}


@router.post("/aml-monitor")
async def audit_aml_monitor(
    file: UploadFile = File(...),
    threshold_amount: float = Form(10000),
    jurisdiction: str = Form(...),
    client_name: str | None = Form(None),
    db: Session = Depends(get_db),
):
    csv_text, row_count = await read_upload_as_csv_text(file)
    user = f"""Jurisdiction: {jurisdiction}
Flagging threshold amount (for context): {threshold_amount}
Transaction data ({row_count} rows):
{csv_text}
"""
    result = invoke_audit_json(SYS_AML, user)
    run = _persist(
        db,
        agent_type="aml-monitor",
        client_name=client_name,
        file_name=file.filename,
        result=result,
    )
    return {"run_id": run.id, "result": result}


@router.get("/runs")
def list_audit_runs(
    agent_type: str | None = None,
    limit: int = 25,
    db: Session = Depends(get_db),
):
    q = db.query(AuditRun).order_by(AuditRun.run_timestamp.desc())
    if agent_type:
        q = q.filter(AuditRun.agent_type == agent_type)
    rows = q.limit(min(limit, 100)).all()
    return {
        "runs": [
            {
                "id": r.id,
                "agent_type": r.agent_type,
                "client_name": r.client_name,
                "file_name": r.file_name,
                "run_timestamp": r.run_timestamp.isoformat() if r.run_timestamp else None,
                "result_summary": r.result_summary,
            }
            for r in rows
        ]
    }


@router.get("/runs/{run_id}/pdf")
def download_audit_pdf(
    run_id: int,
    db: Session = Depends(get_db),
):
    row = db.query(AuditRun).filter(AuditRun.id == run_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="Run not found.")
    labels = {
        "evidence-collector": "Audit Evidence Collection Agent",
        "ifrs-checker": "IFRS Compliance Checker Agent",
        "controls-tester": "Internal Controls Testing Agent",
        "sox-checker": "SOX Compliance Checker Agent",
        "aml-monitor": "AML Transaction Monitor Agent",
    }
    agent_name = labels.get(row.agent_type, row.agent_type)
    result = row.full_result if isinstance(row.full_result, dict) else {}
    pdf = build_audit_pdf_bytes(
        agent_name=agent_name,
        client_name=row.client_name,
        run_at=row.run_timestamp,
        result=result,
    )
    fname = f"audit-{row.agent_type}-{run_id}.pdf"
    return Response(
        content=pdf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@router.get("/runs/{run_id}/evidence-checklist.xlsx")
def download_evidence_checklist_xlsx(
    run_id: int,
    db: Session = Depends(get_db),
):
    row = db.query(AuditRun).filter(AuditRun.id == run_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="Run not found.")
    if row.agent_type != "evidence-collector":
        raise HTTPException(status_code=400, detail="Checklist is only available for evidence collector runs.")
    result = row.full_result if isinstance(row.full_result, dict) else {}
    evidence_rows = result.get("audit_evidence_required") or []

    wb = Workbook()
    ws = wb.active
    ws.title = "Evidence Checklist"
    ws.append(["Priority", "Reference", "Amount", "Risk Level", "Pattern Detected", "Evidence Needed", "ISA Reference", "Status"])
    for item in evidence_rows:
        ws.append(
            [
                item.get("priority", ""),
                item.get("entry_reference", ""),
                item.get("amount", 0),
                item.get("risk_level", ""),
                item.get("pattern_detected", ""),
                item.get("evidence_needed", ""),
                item.get("isa_reference", ""),
                item.get("status", "Pending"),
            ]
        )
    bio = io.BytesIO()
    wb.save(bio)
    fname = f"evidence-checklist-{run_id}.xlsx"
    return Response(
        content=bio.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


@router.get("/runs/{run_id}/management-letter.docx")
def download_management_letter_docx(
    run_id: int,
    db: Session = Depends(get_db),
):
    row = db.query(AuditRun).filter(AuditRun.id == run_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="Run not found.")
    if row.agent_type != "evidence-collector":
        raise HTTPException(status_code=400, detail="Management letter is only available for evidence collector runs.")
    result = row.full_result if isinstance(row.full_result, dict) else {}
    points = result.get("management_letter_points") or []
    findings = result.get("auditor_findings") or ""
    summary = result.get("audit_summary") or {}

    doc = Document()
    doc.add_heading("Management Letter", level=1)
    doc.add_paragraph(f"Client: {row.client_name or 'N/A'}")
    doc.add_paragraph(f"Run date: {row.run_timestamp.isoformat() if row.run_timestamp else 'N/A'}")
    doc.add_paragraph(f"Audit risk rating: {summary.get('audit_risk_rating', 'N/A')}")
    if findings:
        doc.add_heading("Auditor Findings", level=2)
        doc.add_paragraph(str(findings))
    doc.add_heading("Management Letter Points", level=2)
    if points:
        for p in points:
            doc.add_paragraph(str(p), style="List Bullet")
    else:
        doc.add_paragraph("No management letter points generated.")

    bio = io.BytesIO()
    doc.save(bio)
    fname = f"management-letter-{run_id}.docx"
    return Response(
        content=bio.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


class GoingConcernIn(BaseModel):
    workspace_id: str = Field(..., min_length=1)
    company_id: str = Field(..., min_length=1)
    period: str = Field(..., min_length=4)
    financial_data: dict[str, Any] | None = None


def _to_float(v: Any) -> float:
    try:
        return float(v or 0)
    except (TypeError, ValueError):
        return 0.0


def _previous_period(period: str) -> str | None:
    p = (period or "").strip()
    # YYYY-QN
    if len(p) == 7 and p[4] == "-" and p[5].upper() == "Q" and p[6].isdigit():
        year = int(p[:4])
        q = int(p[6])
        if 1 <= q <= 4:
            if q == 1:
                return f"{year - 1}-Q4"
            return f"{year}-Q{q - 1}"
    # YYYY-MM
    if len(p) == 7 and p[4] == "-" and p[5:7].isdigit():
        year = int(p[:4])
        month = int(p[5:7])
        if 1 <= month <= 12:
            if month == 1:
                return f"{year - 1}-12"
            return f"{year}-{month - 1:02d}"
    return None


def _going_concern_level(triggered_count: int) -> str:
    if triggered_count <= 0:
        return "none"
    if triggered_count <= 2:
        return "low"
    if triggered_count <= 4:
        return "medium"
    return "high"


def _is_revenue_account(account_type: str | None, account_code: str | None, account_name: str | None) -> bool:
    t = str(account_type or "").strip().lower()
    if t in {"income", "revenue"}:
        return True
    name = str(account_name or "").lower()
    if "revenue" in name or "income" in name or "sales" in name:
        return True
    code = str(account_code or "").strip()
    return bool(code.startswith("4"))


def _is_expense_account(account_type: str | None, account_code: str | None, account_name: str | None) -> bool:
    t = str(account_type or "").strip().lower()
    if t in {"expense", "expenses", "cost of sales", "cogs"}:
        return True
    name = str(account_name or "").lower()
    if "expense" in name or "cost of" in name or "fees" in name:
        return True
    code = str(account_code or "").strip()
    return bool(code[:1] in {"5", "6", "7"})


@router.post("/going-concern")
def audit_going_concern(
    body: GoingConcernIn,
    db: Session = Depends(get_db),
):
    start, end = parse_period_dates(body.period)

    # 1) Net loss from uae_journal_entries (CR revenue < DR expenses)
    je_line_rows = (
        db.query(
            UAEJournalLine.account_code,
            UAEJournalLine.account_name,
            UAEJournalLine.debit,
            UAEJournalLine.credit,
            UAEAccount.account_type,
        )
        .join(UAEJournalEntry, UAEJournalEntry.id == UAEJournalLine.journal_entry_id)
        .outerjoin(UAEAccount, UAEAccount.id == UAEJournalLine.account_id)
        .filter(
            UAEJournalEntry.tenant_id == body.workspace_id,
            UAEJournalEntry.company_id == body.company_id,
            UAEJournalEntry.entry_date >= start,
            UAEJournalEntry.entry_date <= end,
            UAEJournalEntry.status == "posted",
        )
        .all()
    )
    je_revenue_cr = 0.0
    je_expense_dr = 0.0
    for row in je_line_rows:
        if _is_revenue_account(row.account_type, row.account_code, row.account_name):
            je_revenue_cr += _to_float(row.credit)
        if _is_expense_account(row.account_type, row.account_code, row.account_name):
            je_expense_dr += _to_float(row.debit)
    net_loss_current_period = je_revenue_cr < je_expense_dr

    # GulfTax rows for VAT stress + revenue trend (period window)
    gt_rows = (
        db.query(GulftaxTransaction)
        .filter(
            GulftaxTransaction.tenant_id == body.workspace_id,
            GulftaxTransaction.company_id == body.company_id,
            GulftaxTransaction.transaction_date >= start,
            GulftaxTransaction.transaction_date <= end,
            GulftaxTransaction.status == "posted",
        )
        .all()
    )
    period_revenue = sum(
        max(0.0, _to_float(r.gross_amount) - _to_float(r.vat_amount))
        for r in gt_rows
        if str(r.direction or "").lower() == "output"
    )

    # 2) AR aging stress (>90 days >30% of total AR)
    ar_open_rows = (
        db.query(
            UAESalesInvoice.due_date,
            UAESalesInvoice.outstanding,
            UAESalesInvoice.status,
        )
        .filter(
            UAESalesInvoice.tenant_id == body.workspace_id,
            UAESalesInvoice.company_id == body.company_id,
            UAESalesInvoice.status.in_(("sent", "overdue", "partial", "draft")),
            UAESalesInvoice.outstanding > 0,
        )
        .all()
    )
    total_ar = sum(_to_float(r.outstanding) for r in ar_open_rows)
    overdue_90_ar = 0.0
    for row in ar_open_rows:
        if row.due_date and (end - row.due_date).days > 90:
            overdue_90_ar += _to_float(row.outstanding)
    ar_overdue_ratio = (overdue_90_ar / total_ar) if total_ar > 0 else 0.0
    ar_overdue_above_30pct = ar_overdue_ratio > 0.30

    # 3) AP payment delays (>60 days overdue >25% of total open AP)
    ap_open_rows = (
        db.query(ApInvoice)
        .filter(
            ApInvoice.tenant_id == body.workspace_id,
            ApInvoice.company_id == body.company_id,
        )
        .all()
    )
    total_ap_open = 0.0
    overdue_60_ap = 0.0
    for row in ap_open_rows:
        s = str(row.status or "").lower()
        if s in {"paid", "cancelled", "canceled", "void"}:
            continue
        amt = _to_float(row.total_amount)
        if amt <= 0:
            continue
        total_ap_open += amt
        if row.due_date and (end - row.due_date).days > 60:
            overdue_60_ap += amt
    ap_overdue_ratio = (overdue_60_ap / total_ap_open) if total_ap_open > 0 else 0.0
    ap_payment_delays = ap_overdue_ratio > 0.25

    # 4) VAT/CT compliance stress
    vat_overdue = any(
        str(r.direction or "").lower() == "output"
        and str(r.status or "").lower() not in {"paid", "settled"}
        and (end - r.transaction_date).days > 45
        for r in gt_rows
        if r.transaction_date
    )
    ct_rows = (
        db.query(CtReturn)
        .filter(
            CtReturn.tenant_id == body.workspace_id,
            CtReturn.company_id == body.company_id,
            CtReturn.period_start <= end,
            CtReturn.period_end >= start,
        )
        .order_by(CtReturn.updated_at.desc())
        .all()
    )
    ct_missing_or_overdue = (len(ct_rows) == 0) or all(
        str(r.status or "").lower() not in {"approved", "filed"} for r in ct_rows
    )
    vat_compliance_stress = vat_overdue or ct_missing_or_overdue

    # 5) Revenue decline over 20% vs prior period from gulftax_transactions
    prev_period = _previous_period(body.period)
    prev_revenue = 0.0
    if prev_period:
        prev_rows = (
            db.query(GulftaxTransaction)
            .filter(
                GulftaxTransaction.tenant_id == body.workspace_id,
                GulftaxTransaction.company_id == body.company_id,
                GulftaxTransaction.tax_period == prev_period,
                GulftaxTransaction.direction == "output",
                GulftaxTransaction.status == "posted",
            )
            .all()
        )
        prev_revenue = sum(
            max(0.0, _to_float(r.gross_amount) - _to_float(r.vat_amount)) for r in prev_rows
        )
    revenue_decline_over_20pct = bool(prev_revenue > 0 and period_revenue < (0.8 * prev_revenue))

    # 6) Negative working capital (proxy from open AR vs open AP)
    negative_working_capital = (ap_overdue_ratio > 0.25 and ar_overdue_ratio > 0.30) or (
        total_ap_open > total_ar and total_ap_open > 0
    )

    indicators = {
        "net_loss_current_period": bool(net_loss_current_period),
        "negative_working_capital": bool(negative_working_capital),
        "ar_overdue_above_30pct": bool(ar_overdue_above_30pct),
        "ap_payment_delays": bool(ap_payment_delays),
        "vat_compliance_stress": bool(vat_compliance_stress),
        "revenue_decline_over_20pct": bool(revenue_decline_over_20pct),
    }
    overrides = (body.financial_data or {}).get("overrides") if isinstance(body.financial_data, dict) else None
    if isinstance(overrides, dict):
        for key in list(indicators.keys()):
            if key in overrides:
                indicators[key] = bool(overrides[key])
    triggered_count = sum(1 for v in indicators.values() if v)
    level = _going_concern_level(triggered_count)

    user_prompt = (
        f"Company ID: {body.company_id}\n"
        f"Workspace ID: {body.workspace_id}\n"
        f"Period: {body.period}\n"
        f"Triggered indicators: {triggered_count} of 6\n"
        f"Indicators: {json.dumps(indicators)}\n"
        f"Metrics: {{"
        f"\"je_revenue_credits\": {round(je_revenue_cr, 2)}, "
        f"\"je_expense_debits\": {round(je_expense_dr, 2)}, "
        f"\"gulftax_period_revenue\": {round(period_revenue, 2)}, "
        f"\"ar_overdue_ratio\": {round(ar_overdue_ratio, 4)}, "
        f"\"ap_overdue_ratio\": {round(ap_overdue_ratio, 4)}, "
        f"\"total_ar\": {round(total_ar, 2)}, "
        f"\"total_ap_open\": {round(total_ap_open, 2)}, "
        f"\"previous_period\": \"{prev_period}\", "
        f"\"previous_period_revenue\": {round(prev_revenue, 2)}"
        f"}}"
    )
    try:
        narrative = invoke(
            prompt=user_prompt,
            system=GOING_CONCERN_SYSTEM_PROMPT,
            max_tokens=500,
            temperature=0.2,
            model_id="claude-sonnet-4-6",
        ).strip()
    except Exception as exc:
        # Keep endpoint live even if LLM is down.
        narrative = (
            f"Going concern assessment generated with deterministic indicators only: "
            f"{triggered_count}/6 indicators triggered ({level}). AI narrative unavailable ({exc})."
        )

    generated_at = datetime.now(timezone.utc).isoformat()
    result = {
        "period": body.period,
        "workspace_id": body.workspace_id,
        "company_id": body.company_id,
        "indicators": indicators,
        "triggered_count": triggered_count,
        "going_concern_level": level,
        "narrative": narrative,
        "generated_at": generated_at,
        "isa_570_reference": "ISA 570",
    }
    run = _persist(
        db,
        agent_type="going_concern",
        client_name=body.company_id,
        file_name=None,
        result=result,
    )
    return {"run_id": run.id, **result}
